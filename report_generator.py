"""report_generator.py — Word report generation"""
import os
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _sf(run, size=10, bold=False, color=None, italic=False):
    run.font.name='Calibri'; run.font.size=Pt(size)
    run.font.bold=bold; run.font.italic=italic
    if color: run.font.color.rgb=RGBColor(*color)

def _p(doc, text="", size=10, bold=False, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4, italic=False):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if text:
        r=p.add_run(text); _sf(r,size,bold,color,italic)
    return p

def _hr(doc, color="1F4E79"):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(2)
    p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),color)
    pBdr.append(bot); pPr.append(pBdr)

def _shade(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),hex_color); tcPr.append(shd)

def generate_word_report(stats: dict, report_text: str) -> bytes:
    doc=Document()
    for s in doc.sections:
        s.top_margin=Cm(2); s.bottom_margin=Cm(2)
        s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    _p(doc,"Y K JONEJA & CO.",14,True,(31,78,121),WD_ALIGN_PARAGRAPH.CENTER,sa=2)
    _p(doc,"Advocates & Tax Consultants  |  Faridabad",
       9,False,(89,89,89),WD_ALIGN_PARAGRAPH.CENTER,italic=True,sa=2)
    _hr(doc); _p(doc,"")
    _p(doc,"GSTR-1 FILING READINESS REPORT",13,True,(31,78,121),
       WD_ALIGN_PARAGRAPH.CENTER,sa=2)
    _p(doc,f"Generated: {datetime.now().strftime('%d %B %Y  %H:%M')}",
       8,False,(127,127,127),WD_ALIGN_PARAGRAPH.CENTER,italic=True,sa=8)
    if stats:
        tbl=doc.add_table(rows=2,cols=4); tbl.style="Table Grid"
        tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        for r_idx,rd in enumerate([
            ("Client",stats.get("client_name",""),"GSTIN",stats.get("gstin","")),
            ("Period",stats.get("tax_period",""),"Date",datetime.now().strftime("%d-%b-%Y"))
        ]):
            row=tbl.rows[r_idx]
            for ci,(lb,vl) in enumerate([(rd[0],rd[1]),(rd[2],rd[3])]):
                lc,vc=row.cells[ci*2],row.cells[ci*2+1]
                _shade(lc,"DCE6F1")
                lr=lc.paragraphs[0].add_run(lb); _sf(lr,9,True,(31,78,121))
                vr=vc.paragraphs[0].add_run(vl); _sf(vr,9)
        _p(doc,"")
        rate=stats.get("match_rate_pct",0)
        verdict=("Ready to file" if rate>=95 else
                 "Review required" if rate>=85 else "Hold — do not file")
        vbg=("E2EFDA" if "Ready" in verdict else
             "FFF2CC" if "Review" in verdict else "FCE4D6")
        vc2=("0,97,0" if "Ready" in verdict else
             "197,90,17" if "Review" in verdict else "192,0,0")
        vtbl=doc.add_table(rows=1,cols=1); vtbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        vc3=vtbl.rows[0].cells[0]; _shade(vc3,vbg)
        vp=vc3.paragraphs[0]; vp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        vrun=vp.add_run(f"FILING VERDICT:  {verdict.upper()}")
        rgb=[int(x) for x in vc2.split(",")]
        _sf(vrun,13,True,tuple(rgb)); _p(doc,"")
        _p(doc,"RECONCILIATION STATISTICS",10,True,(31,78,121),sb=4,sa=4); _hr(doc)
        intra=stats.get("intra_state_threshold",50000)
        if isinstance(intra,str): intra=50000
        rows=[
            ("Total Sales Invoices",str(stats.get("total_sales",""))),
            ("Total EWBs",str(stats.get("total_ewb",""))),
            ("Delivery Challans",str(stats.get("delivery_challans",""))),
            ("Matched",str(stats.get("matched_ok",""))),
            ("Value Mismatches",str(stats.get("val_mismatch",""))),
            ("Sales No EWB",str(stats.get("sales_no_ewb",""))),
            ("EWB No Match",str(stats.get("ewb_no_sales",""))),
            ("E-Invoice IRN Missing",str(stats.get("einv_missing_irn",""))),
            ("Match Rate %",f"{rate}%"),
            ("Total Issues",str(stats.get("issues_count",""))),
        ]
        stbl=doc.add_table(rows=len(rows),cols=2); stbl.style="Table Grid"
        for i,(lb,vl) in enumerate(rows):
            lc2,vc4=stbl.rows[i].cells[0],stbl.rows[i].cells[1]
            _shade(lc2,"DCE6F1"); _shade(vc4,"FFFFFF")
            lr2=lc2.paragraphs[0].add_run(lb); _sf(lr2,9,True,(31,78,121))
            vc4.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            vr2=vc4.paragraphs[0].add_run(vl); _sf(vr2,9)
        _p(doc,"")
    _p(doc,"DETAILED ANALYSIS",10,True,(31,78,121),sb=6,sa=3); _hr(doc)
    for line in report_text.split("\n"):
        s2=line.strip()
        if not s2: _p(doc,"")
        elif s2.startswith("=") or s2.startswith("-"): _hr(doc)
        elif (any(s2.startswith(f"{i}.") for i in range(1,10)) or
              s2.upper().startswith("FILING VERDICT") or
              s2.upper().startswith("GSTR-1")):
            _p(doc,s2,10,True,(31,78,121),sb=4,sa=2)
        else:
            _p(doc,s2,9,sa=2)
    _hr(doc)
    _p(doc,"Y K Joneja & Co.  |  Advocates & Tax Consultants  |  Faridabad  |  Confidential",
       8,False,(127,127,127),WD_ALIGN_PARAGRAPH.CENTER,italic=True,sb=4)
    buf=BytesIO(); doc.save(buf); return buf.getvalue()
